from docx import Document 
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Inches, Cm

document = Document()
style = document.styles['Normal']
style.paragraph_format.line_spacing = 0.73
font = style.font
font.name = 'Myriad Pro'
font.size = Pt(16)
font.bold = True

section = document.sections[0]

sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')

sections = document.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.4478)
    section.left_margin = Cm(0.635)
    section.right_margin = Cm(0.7874)

# original top margin value = 2.1082
# turn adding label chunks into a function

p = document.add_paragraph()
firstLine = p.add_run("UDI COMPLIANT")
firstLine.bold = True
firstLine.underline = True
secondLine = p.add_run(": DS-SUM-500")
secondLine.bold = True
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph()
firstLine = p.add_run("Mouthpiece – Medium (20-Pack)")
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph()
firstLine = p.add_run("LOT 006-SUM-06")
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER



document.save("Outer Carton Labels Test.docx")